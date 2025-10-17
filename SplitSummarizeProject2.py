import os 
import streamlit as st
from dotenv import load_dotenv, find_dotenv
from langchain_openai import OpenAI
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.summarize import load_summarize_chain
import pandas as pd
import io 
from io import BytesIO
from bs4 import BeautifulSoup
import chardet  
from pypdf import PdfReader
from langchain_core.documents import Document
import time 
from docx import Document as DocxDocument
from openai import RateLimitError, APIError

_ = load_dotenv(find_dotenv())
openai_api_key = os.environ["OPENAI_API_KEY"]   

def load_LLM(openai_api_key):
    llm = OpenAI(temperature = 0, openai_api_key=openai_api_key)
    return llm

def read_uploaded_file(uploaded_file) -> str | None:
    """
    Read uploaded_file (Streamlit UploadedFile) and return plain text.
    Supports: .txt, .pdf, .html/.htm, .mhtml/.mht (basic).
    Handles Edge's 'multipart/related' MHTML saves.
    """
    if uploaded_file is None:
        return None
    # raw bytes + metadata
    bytes_data = uploaded_file.getvalue()
    name = (getattr(uploaded_file, "name", "") or "").lower()
    mime = (getattr(uploaded_file, "type", "") or "").lower()

    # --- PDF ---
    if name.endswith(".pdf") or mime == "application/pdf":
        try:
            reader = PdfReader(BytesIO(bytes_data))
            pages = []
            for p in reader.pages:
                txt = p.extract_text() or ""
                pages.append(txt)
            return "\n\n".join(pages).strip()
        except Exception as e:
            st.error(f"PDF read error: {e}")
            return None
    
    # ---docx ----
    
    if name.endswith(".docx") or \
       mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            doc = DocxDocument(BytesIO(bytes_data))
            parts = []
            # paragraphs
            parts.extend(p.text for p in doc.paragraphs if p.text)
            # tables (optional but useful)
            for table in doc.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells if cell.text))
            return "\n".join(parts).strip()
        except Exception as e:
            st.error(f"DOCX read error: {e}")
            return None 

    # --- HTML / MHTML ---
    if (
        name.endswith((".html", ".htm", ".mhtml", ".mht"))
        or "html" in mime
        or "multipart/related" in mime  # Edge's MHTML
    ):
        try:
            # First try UTF-8, then fallback to detected encoding
            s = bytes_data.decode("utf-8", errors="replace")
            if ("<html" not in s.lower()) and ("<!doctype" not in s.lower()):
                enc = chardet.detect(bytes_data).get("encoding") or "utf-8"
                s = bytes_data.decode(enc, errors="replace")

            soup = BeautifulSoup(s, "lxml")
            plain = soup.get_text(separator="\n")
            return plain.strip()
        except Exception as e:
            st.error(f"HTML read error: {e}")
            return None
    # --- Plain text (default) ---
    try:
        enc = chardet.detect(bytes_data).get("encoding") or "utf-8"
        return bytes_data.decode(enc, errors="replace")
    except Exception as e:
        st.error(f"Text decode error: {e}")
        return None


st.set_page_config(page_title="AI Long Text Summarizer")
st.header("AI Long Text Summarizer")

col1, col2 = st.columns(2) 

with col1:
    st.markdown("ChatGPT cannot summarize long texts. Now you can do it with this app")
    
with col2:
    st.write("Contact Tariq Mansaray tariq.mansaray@gmail.com for your AI solutions")
    
st.markdown("## Enter Your OpenAI API Key")

def get_openai_api_key():
    input_text = st.text_input(label= "OpenAI API Key", placeholder = "Ex: sk-2twmA8tfCb8un4...", key = "openai_api_key_input", type = "password")
    return input_text

openai_api_key = get_openai_api_key()

st.markdown("## Upload the text file you want to summarize")

uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "html", "htm", "mhtml","docx"])


st.markdown("### Here is your summary")

if uploaded_file is not None:
    #bytes_data = uploaded_file.getvalue()
    
    #stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
    
    #string_data = stringio.read()
    
    #file_input = string_data
    if not openai_api_key:   
        st.write("Please insert OpenAI API Key. Instructions [here](https://help.openai.com/en/articles/4936850-where-do-i-find-my-secret-api-key)",
                 icon = "⚠️")
        st.stop() 
    
    file_text = read_uploaded_file(uploaded_file)
    if not file_text:
        st.stop()
   
    
    if len(file_text.split(" ")) > 90000:
        st.write("Please enter a shorter file. Maximum length is 90000 words. ")
        st.stop()
        

        
    text_splitter = RecursiveCharacterTextSplitter(
        separators = ["\n\n", "\n"],
        chunk_size = 5000,
        chunk_overlap = 350
    )
    
    docs = [Document(page_content=file_text)]
    splitted_documents = text_splitter.split_documents(docs)
    
    llm = load_LLM(openai_api_key = openai_api_key)
    
    summarize_chain = load_summarize_chain(
        llm=llm,
        chain_type= "map_reduce"
    )
    
    summary_output = summarize_chain.run(splitted_documents) 
    st.write(summary_output)